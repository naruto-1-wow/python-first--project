#start recalling knowledges of python
'''
name = 'Jamila'
age = 18
pi = 3.14
cars = ['bmw','mercedes','range rover']

print(name)
print(age)
print(pi)
print(cars)
'''

'''------------------------------Split Line----------------------------------------------'''
#recall knowledges of string
'''
first_name = 'jamila'
surname = 'smith'
full_name = first_name +' '+ surname


print(first_name)
print(surname)
print(full_name)
print(full_name.capitalize())
print(len(first_name))
print(len(surname))
'''

'''-----------------------------Split Line----------------------------------------------------------'''
#recall knowledges of numbers
'''
addition = 10 + 5
subtraction = 10 - 5
multiplication = 10 * 5
division = 10 / 5
mod = 10 % 5#the remainder of 10/5


print(addition)
print(subtraction)
print(multiplication)
print(division)
print(mod)
'''

'''------------------------------------------Split Line----------------------------------------------'''
#recall knowledges of boolean

from operator import truediv
from re import L

'''
print(10 < 10)
print(0 == 0)
print(18 > 5)
print('jamila'.endswith('s'))
print('jamila'.endswith('a'))
is_adult = True
is_teenager = False
'''

'''--------------------------------------------Split Line------------------------------------------------'''
#recall knowledges of ifStatement
'''
is_adult = True
is_teenager = False
age = 18

if is_adult:
    print('is adult')
if is_teenager:
    print('is teenager')

if age >= 18:#if .... else  ...  statements
    print('adult')
else:
    print('not an adult')
'''

'''--------------------------------------------Split Line------------------------------------------------'''
#recall knowledges of list 
'''
cars = ['bmw','tesla','mercedes']
print(len(cars))
print(cars)
print(cars[0])#the index starts from 0
#cars[3]-------------------------------IndexErrer:list index out of range
'''


'''------------------------------------------Split Line--------------------------------------------------'''
#recall knowledges of loops
'''
cars = ['bmw','tesla','mercedes','toyota','honda']

for car in cars:
    print(car)

for car in cars:
    print(car.capitalize())#don't forget "()" behind the method "capitalize()"

for car in cars:
    if car == 'bmw':
        print(car.upper())#don't forget "()" behind the method "upper()"
    else:
        print(car.capitalize())    
    '''


'''------------------------------------------Split Line-----------------------------------------------'''
#recall knowledges of while loops
'''
number = 0

while number <= 10:
    print(number)
    number = number + 1
else:
    print('while loop ended and value of number is' + ' '+str(number))
'''

'''============================================Split Line==========================================='''
#learn the knowledge of functions
'''
age = 18
age2 = 17

def check_age(age):
    if age < 18:
       print('not an adult')
    else:
       print('an adult')    

check_age(age)
check_age(age2)
check_age(18)
check_age(17)
'''
#P.S:Watch out for the space. It's different from the space in c++,jave... it has meaning in syntax

'''============================================Split Line========================================'''
#learn the knowledges of class
'''
class Person:#
    def __init__(self,name,age):#This function defines properties
        self.name = name
        self.age = age

john = Person('John',22)
mariam = Person('Mariam',18)

print(john.name + ' ' +str(john.age))
print(mariam.name + ' ' + str(mariam.age))
'''

'''=================================================Split Line=============================================='''
#learn classes and behavior and objects
'''
class Person:
    def __init__(self,name,age):#properties
        self.name = name
        self.age = age
    
    def walk(self):#behivior
        print(self.name + ' is walking...')
    
    def speak(self):#behivior
        print('Hello, my name is' + self.name + ' ' + 'and i am' + ' ' + str(self.age) + ' ' + 'years old')


john = Person('John',22)
mariam = Person('Mariam',18)

print(john.name + ' ' +str(john.age))
john.speak()
john.walk()

print(mariam.name + ' ' + str(mariam.age))
mariam.speak()
mariam.walk()
'''

'''=================================================Split Line==============================================='''
#A project named "cv"

from docx import Document# With this statement, We are importing the "Document" from the "docx" we just installed
from docx.shared import Inches# With this statement, We can set the size of picture
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()# Adding "()" means invoking this fun8.0))
   
# profile picture
document.add_picture(
    'Naruto.jfif',
    width = Inches(4.0)
    )

name = input('What is your name ?')
speak('Hello' + name + 'how are you today ? ')

speak('What is your phone number ?')
phone_number = input('What is your phone number ?')
email_qq = input('What is your qq email ?')
email_g = input('What is your google email ?')

#add context about me into the docx
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email_qq + ' | ' + email_g
)

# about me
document.add_heading('About me:')
document.add_paragraph(input('Tell me about your self?'))

#work experience
document.add_heading('Work experience:')
p = document.add_paragraph()

company = input('What company are you used to work in?')
from_date = input('When did you start to work?')
to_date = input('When did your job finished?')
# add text to existing paragraphs
p.add_run(company + ' ').bold = True #bold:加粗字体
p.add_run(from_date + ' ' + ' to ' + to_date +'\n').italic = True #italic:斜体字

experience_details = input('Describe your experience at ' + company + ':')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input (
        'Do you have more experiences? Yes or no ?'
    )
    if has_more_experiences.lower() == 'yes':
        document.add_heading('Work experience:')
        p = document.add_paragraph()

        company = input('What company are you used to work in?')
        from_date = input('When did you start to work?')
        to_date = input('When did your job finished?')

        p.add_run(company + ' ').bold = True #bold:加粗字体
        p.add_run(from_date + ' ' + ' to ' + to_date +'\n').italic = True #italic:斜体字
        experience_details = input('Describe your experience at ' + company + ':')
        p.add_run(experience_details)
    else:
        break

#skills
document.add_heading('Skills:')
skill = input('Enter skill :')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Any more akills ? yes or no ?')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill :')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using project'

#save the docx which we just modified
document.save('cv.docx')#create a docx document named "cv"

'''                           What i've learned  
    1.pip3 is a package management system which i can use it to download the package i need
    2.package "docx" allows me to modify a docx file. It provides me with many useful functions
      such as:add_paragraph(),add_picture()
    3.make good use of file called Requirements.txt . And then we are going to have all of the packages this project needs.
...
'''

