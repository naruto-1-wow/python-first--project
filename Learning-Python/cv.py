#A project named "cv"
''''''
from docx import Document# With this statement, We are importing the "Document" from the "docx" we just installed
from docx.shared import Inches# With this statement, We can set the size of picture

document = Document()# Adding "()" means invoking this fun8.0))

# profile picture
document.add_picture(
    'Naruto.jfif',
    width = Inches(4.0)
    )

name = input('What is your name ?')
phone_number = input('What is your phone number ?')
email_qq = input('What is your qq email ?')
email_g = input('What is your google email ?')

#add context about me into the docx
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email_qq + ' | ' + email_g
)

# about me
document.add_heading('About me:')
document.add_paragraph(input('Tell me about your self?'))

#work experience
document.add_heading('Work experience:')
p = document.add_paragraph()

company = input('What company are you used to work in?')
from_date = input('When did you start to work?')
to_date = input('When did your job finished?')
# add text to existing paragraphs
p.add_run(company + ' ').bold = True #bold:加粗字体
p.add_run(from_date + ' ' + ' to ' + to_date +'\n').italic = True #italic:斜体字

experience_details = input('Describe your experience at ' + company + ':')
p.add_run(experience_details)

# more experiences
def new_func(document, skill):
    p = document.add_paragraph(skill)
    return p

while True:
    has_more_experiences = input ('Do you have more experiences? Yes or no ?')
    if has_more_experiences.lower() == 'yes':
        document.add_heading('Work experience:')
        p = document.add_paragraph()

        company = input('What company are you used to work in?')
        from_date = input('When did you start to work?')
        to_date = input('When did your job finished?')

        # add text to existing paragraphs
        p.add_run(company + ' ').bold = True #bold:加粗字体
        p.add_run(from_date + ' ' + ' to ' + to_date +'\n').italic = True #italic:斜体字
        experience_details = input('Describe your experience at ' + company + ':')
        p.add_run(experience_details)
    else:
        break

    # more skills
document.add_heading('Skills:')
skill = input('Enter skill :')
p = new_func(document, skill)
p.style = 'List Bullet'

while True:
     has_more_skills = input('Any more akills ? yes or no ?')
     if has_more_skills.lower() == 'yes':
        skill = input('Enter skill :')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
     else:
        break


    

#save the docx which we just modified
document.save('cv.docx')#create a docx document named "cv"


'''                           What i've learned  
    1.pip3 is a package management system which i can use it to download the package i need
    2.package "docx" allows me to modify a docx file. It provides me with many useful functions
      such as:add_paragraph(),add_picture()...
'''

